from telegram import Update, ReplyKeyboardMarkup
from telegram.ext import CommandHandler, CallbackContext
from .model import Model
from docx import Document
import io

TEACHER = 1
Teacher = Model('teacher')
Parent = Model('parent')


def get_result(update: Update, context: CallbackContext):
    username = update.message.from_user.username

    parents = Parent.find_all({"teacher_username": username})

    document = Document()

    document.add_heading('Отчет посещаемости', 0)

    table = document.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ФИО ученика'
    hdr_cells[1].text = 'Идет в школу?'
    hdr_cells[2].text = 'Причина'

    for parent in parents:
        row_cells = table.add_row().cells
        row_cells[0].text = parent["children_fio"]
        if parent["is_attend"] == 1:
            row_cells[1].text = "Да"
            row_cells[2].text = ""
        elif parent["is_attend"] == 2:
            row_cells[1].text = "Неизвестно"
            row_cells[2].text = ""
        elif parent["is_attend"] == 0:
            row_cells[1].text = "Нет"
            row_cells[2].text = parent["reason"]
    print(1)
    document.add_page_break()

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    print(2)
    context.bot.send_document(
        chat_id=update.message.chat_id,
        document=file_stream,
        filename="result.docx"
    )


def start_quest(update: Update, context: CallbackContext):
    username = update.message.from_user.username

    parents = Parent.find_all({"teacher_username": username})

    buttons = [
        ['Да', 'Нет'],
    ]
    keyboard = ReplyKeyboardMarkup(buttons, one_time_keyboard=True)

    for parent in parents:
        Parent.update(
            {"username": username},
            {"$set": {"is_attend": 2, "reason": ""}}
        )
        context.bot.send_message(
            chat_id=parent['chat_id'],
            text="Ваш ребенок пойдет в школу?",
            reply_markup=keyboard
        )

    context.bot.send_message(
        chat_id=update.message.chat_id,
        text="Сообщения родителям успешно разосланы."
    )


teacher_conv = [
    CommandHandler('get_result', get_result),
    CommandHandler('start_quest', start_quest),
]
